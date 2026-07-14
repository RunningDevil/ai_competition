from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any, Dict, List, Tuple
from xml.etree import ElementTree as ET

from comment_parser import looks_like_comment, parse_comment_text
from office_common import (
    convert_legacy_office,
    extract_printable_text_from_binary,
    filter_comments,
    fixed_relative_path,
    get_extension,
    make_error,
    normalize_path_text,
    replacement_from_comment,
    source_relative_to_docs,
)


WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}
W_P = f"{{{WORD_NS['w']}}}p"
W_T = f"{{{WORD_NS['w']}}}t"


def _text_from_xml_element(element: ET.Element) -> str:
    parts = [node.text or "" for node in element.findall(".//w:t", WORD_NS)]
    text = "".join(parts).strip()
    if text:
        return text
    return "".join(part.strip() for part in element.itertext() if part and part.strip()).strip()


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _comment_elements(root: ET.Element) -> List[ET.Element]:
    nodes = [node for node in root.iter() if _local_name(node.tag) in {"comment", "modernComment", "cm"}]
    return nodes or [root]


def _extract_docx_text(path: Path, source_rel: str) -> List[Dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    texts: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            texts.append({"source": source_rel, "file_type": "docx", "location": f"paragraph:{index}", "text": text})
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append({
                    "source": source_rel,
                    "file_type": "docx",
                    "location": f"table:{table_index}:row:{row_index}",
                    "text": "\t".join(cells),
                })
    return texts


def _extract_docx_comments(path: Path, source_rel: str) -> List[Dict[str, Any]]:
    comments: List[Dict[str, Any]] = []
    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
        comment_files = [name for name in names if name.startswith("word/") and "comment" in name.lower() and name.endswith(".xml")]
        for comment_file in comment_files:
            root = ET.fromstring(package.read(comment_file))
            for index, comment in enumerate(_comment_elements(root), start=1):
                text = _text_from_xml_element(comment)
                if text:
                    comments.append(parse_comment_text(text, source_rel, "docx", f"{comment_file}:comment:{index}"))
    return comments


def extract_text(path: Path, wiki_root: str | Path, logs: List[str]) -> List[Dict[str, Any]]:
    ext = get_extension(path)
    source_rel = source_relative_to_docs(path, wiki_root)
    if ext == ".docx":
        return _extract_docx_text(path, source_rel)
    if ext == ".doc":
        converted = convert_legacy_office(path, logs)
        if converted and converted.suffix.lower() == ".docx":
            return _extract_docx_text(converted, source_rel)
        fallback = extract_printable_text_from_binary(path)
        return [{"source": source_rel, "file_type": "doc", "location": "legacy:fallback_text", "text": fallback}] if fallback else []
    return []


def extract_comments(path: Path, wiki_root: str | Path, logs: List[str]) -> List[Dict[str, Any]]:
    ext = get_extension(path)
    source_rel = source_relative_to_docs(path, wiki_root)
    if ext == ".docx":
        return _extract_docx_comments(path, source_rel)
    if ext == ".doc":
        converted = convert_legacy_office(path, logs)
        if converted and converted.suffix.lower() == ".docx":
            return _extract_docx_comments(converted, source_rel)
        fallback = extract_printable_text_from_binary(path)
        if looks_like_comment(fallback):
            return [parse_comment_text(fallback, source_rel, "doc", "legacy:fallback_text")]
        logs.append("Legacy .doc comments could not be reliably extracted")
    return []


def _replace_in_document_xml(xml_bytes: bytes, old: str, new: str) -> Tuple[bytes, int]:
    root = ET.fromstring(xml_bytes)
    replaced_count = 0
    for paragraph in root.iter(W_P):
        text_nodes = [node for node in paragraph.iter(W_T)]
        if not text_nodes:
            continue
        combined = "".join(node.text or "" for node in text_nodes)
        if old not in combined:
            continue
        updated = combined.replace(old, new, 1)
        text_nodes[0].text = updated
        for node in text_nodes[1:]:
            node.text = ""
        replaced_count += 1
        break
    return ET.tostring(root, encoding="utf-8", xml_declaration=True), replaced_count


def _write_replaced_docx(source: Path, target: Path, old: str, new: str, logs: List[str]) -> int:
    target.parent.mkdir(parents=True, exist_ok=True)
    replaced_count = 0
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data, replaced_count = _replace_in_document_xml(data, old, new)
            zout.writestr(item, data)
    if replaced_count:
        logs.append(f"Replaced Word text: {old} -> {new}")
    else:
        logs.append(f"Replacement text not found in Word body: {old}")
    return replaced_count


def conservative_fix(path: Path, wiki_root: str | Path, logs: List[str], filters: Dict[str, Any] | None = None) -> Dict[str, Any]:
    comments = extract_comments(path, wiki_root, logs)
    if filters:
        comments = filter_comments(comments, filters)
    if not comments:
        return make_error("fix_comments", f"No reliable Word comments found for {normalize_path_text(str(path))}", logs)
    source_rel = source_relative_to_docs(path, wiki_root)
    target_rel = fixed_relative_path(source_rel)
    target_abs = Path(wiki_root) / target_rel
    for comment in comments:
        replacement = replacement_from_comment(comment)
        if not replacement:
            continue
        old, new = replacement
        replaced_count = _write_replaced_docx(path, target_abs, old, new, logs)
        if replaced_count:
            return {
                "status": "ok",
                "task_type": "fix_comments",
                "texts": [],
                "comments": comments,
                "answer": {"source": source_rel, "target": target_rel},
                "fixed_files": [target_rel],
                "logs": logs,
            }
    logs.append("Word comments found, but no deterministic replacement instruction could be applied")
    return make_error("fix_comments", f"No reliable Word repair rule implemented for {normalize_path_text(str(path))}", logs)
